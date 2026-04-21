from io import BytesIO
import uuid
import urllib.request

import pandas as pd
from docx import Document


def build_sample_docx_bytes() -> bytes:
    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("CÔNG TY TNHH TMDV QUẢNG CÁO - IN ẤN QUANG MINH")
    doc.add_paragraph("Số ĐKKD/MST\t3502579019")
    doc.add_paragraph("Ngày cấp\t03/04/2026")
    doc.add_paragraph("Ngày hoạt động\t03/04/2026")
    doc.add_paragraph("Tình trạng\tĐang Hoạt Động")
    doc.add_paragraph(
        "Địa chỉ\tSố 43/9 Tôn Thất Thuyết, Phường Tam Thắng, Thành Phố Vũng Tàu, Bà Rịa - Vũng Tàu"
    )
    doc.add_paragraph("Người đại diện\tNgô Quang Hiển")
    doc.add_paragraph("Điện thoại 0918004372")
    doc.save(buf)
    return buf.getvalue()


def build_multipart(field_name: str, filename: str, content_type: str, file_bytes: bytes):
    boundary = "----Boundary" + uuid.uuid4().hex
    crlf = "\r\n"

    head = (
        f"--{boundary}{crlf}"
        f'Content-Disposition: form-data; name="{field_name}"; filename="{filename}"{crlf}'
        f"Content-Type: {content_type}{crlf}"
        f"{crlf}"
    ).encode("utf-8")

    tail = f"{crlf}--{boundary}--{crlf}".encode("utf-8")
    body = head + file_bytes + tail
    content_type_header = f"multipart/form-data; boundary={boundary}"
    return body, content_type_header


def main():
    docx_bytes = build_sample_docx_bytes()
    body, ct = build_multipart(
        field_name="word_file",
        filename="sample.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_bytes=docx_bytes,
    )

    req = urllib.request.Request(
        "http://127.0.0.1:5000/word-to-excel",
        data=body,
        method="POST",
        headers={
            "Content-Type": ct,
            "Content-Length": str(len(body)),
        },
    )

    with urllib.request.urlopen(req, timeout=30) as resp:
        content = resp.read()
        print("HTTP", resp.status, "bytes", len(content))

    df = pd.read_excel(BytesIO(content), dtype=str)
    print(df[["Khách hàng", "Tên liên hệ", "Điện thoại", "Địa chỉ"]].to_string(index=False))


if __name__ == "__main__":
    main()

