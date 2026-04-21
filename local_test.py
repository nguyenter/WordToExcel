from io import BytesIO

import pandas as pd
from docx import Document

from app.services.word_to_excel_service import convert_docx_to_excel_bytes


def build_sample_docx_bytes() -> BytesIO:
    doc = Document()
    doc.add_paragraph("CÔNG TY TNHH TMDV QUẢNG CÁO - IN ẤN QUANG MINH")
    doc.add_paragraph("Số ĐKKD/MST\t3502579019")
    doc.add_paragraph("Ngày cấp\t03/04/2026")
    doc.add_paragraph("Ngày hoạt động\t03/04/2026")
    doc.add_paragraph("Tình trạng\tĐang Hoạt Động")
    doc.add_paragraph(
        "Địa chỉ\tSố 43/9 Tôn Thất Thuyết, Phường Tam Thắng, Thành Phố Vũng Tàu, Bà Rịa - Vũng Tàu"
    )
    doc.add_paragraph("Người đại diện\tNgô Quang Hiển")
    doc.add_paragraph("Điện thoại 0918004372")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def main():
    docx_bytes = build_sample_docx_bytes()
    xlsx_bytes = convert_docx_to_excel_bytes(docx_bytes)
    df = pd.read_excel(xlsx_bytes)

    # Show a compact view of the parsed fields
    cols = [c for c in ["Khách hàng", "Tên liên hệ", "Điện thoại", "Địa chỉ"] if c in df.columns]
    print(df[cols].to_string(index=False))


if __name__ == "__main__":
    main()

